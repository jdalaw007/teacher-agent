from fastapi import APIRouter, HTTPException, Request, UploadFile, File, Form
import json
import hashlib
import uuid
from pathlib import Path
from app.services.token_store import get_user_id_from_token
from app.services.database import get_db

router = APIRouter()

DATA_DIR = Path(__file__).parent.parent.parent / "data" / "corpus"
DATA_DIR.mkdir(parents=True, exist_ok=True)


def _get_user_id(request: Request) -> str:
    auth = request.headers.get("Authorization", "")
    if auth.startswith("Bearer "):
        return get_user_id_from_token(auth[7:])
    raise HTTPException(401, "Missing authorization")


def _get_class_dir(user_id: str, class_id: str) -> Path:
    d = DATA_DIR / user_id / class_id
    d.mkdir(parents=True, exist_ok=True)
    return d


def _get_index(user_id: str, class_id: str) -> dict:
    idx_file = _get_class_dir(user_id, class_id) / "index.json"
    if idx_file.exists():
        return json.loads(idx_file.read_text())
    return {"documents": {}}


def _save_index(user_id: str, class_id: str, index: dict):
    (_get_class_dir(user_id, class_id) / "index.json").write_text(json.dumps(index, indent=2))


# --- Folders ---

@router.get("/folders")
async def list_folders(request: Request):
    user_id = _get_user_id(request)
    db = get_db()
    try:
        rows = db.execute(
            "SELECT id, class_id, name, parent_id, created_at FROM file_folders WHERE user_id = ? ORDER BY class_id, name",
            (user_id,)
        ).fetchall()
        return {"folders": [dict(r) for r in rows]}
    finally:
        db.close()


@router.post("/folders")
async def create_folder(request: Request):
    user_id = _get_user_id(request)
    body = await request.json()
    name = body.get("name", "").strip()
    class_id = body.get("class_id", "general")
    parent_id = body.get("parent_id")
    if not name:
        raise HTTPException(400, "Folder name required")
    folder_id = str(uuid.uuid4())[:12]
    db = get_db()
    try:
        db.execute(
            "INSERT INTO file_folders (id, user_id, class_id, name, parent_id) VALUES (?, ?, ?, ?, ?)",
            (folder_id, user_id, class_id, name, parent_id)
        )
        db.commit()
        return {"id": folder_id, "class_id": class_id, "name": name, "parent_id": parent_id}
    finally:
        db.close()


@router.delete("/folders/{folder_id}")
async def delete_folder(folder_id: str, request: Request):
    user_id = _get_user_id(request)
    db = get_db()
    try:
        db.execute("DELETE FROM file_folders WHERE id = ? AND user_id = ?", (folder_id, user_id))
        db.commit()
    finally:
        db.close()
    return {"deleted": True}


# --- Documents ---

@router.get("/documents")
async def list_documents(request: Request, class_id: str):
    user_id = _get_user_id(request)
    index = _get_index(user_id, class_id)
    docs = [{"id": doc_id, **info} for doc_id, info in index["documents"].items()]
    return {"documents": docs, "class_id": class_id}


@router.post("/upload")
async def upload_file(
    request: Request,
    file: UploadFile = File(...),
    class_id: str = Form("general"),
    folder_id: str = Form(None),
    title: str = Form(None),
):
    user_id = _get_user_id(request)
    class_dir = _get_class_dir(user_id, class_id)
    content = await file.read()
    filename = file.filename or "untitled"
    doc_title = title or filename

    if filename.lower().endswith('.pdf'):
        try:
            from pypdf import PdfReader
            import io
            pdf = PdfReader(io.BytesIO(content))
            if pdf.is_encrypted:
                if pdf.decrypt("") == 0:
                    raise HTTPException(400, "PDF is password-protected")
            text = "\n".join(page.extract_text() or "" for page in pdf.pages)
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(400, f"Failed to parse PDF: {e}")
    elif filename.lower().endswith('.docx'):
        try:
            from docx import Document
            import io
            doc = Document(io.BytesIO(content))
            text = "\n".join(para.text for para in doc.paragraphs)
        except Exception as e:
            raise HTTPException(400, f"Failed to parse Word document: {e}")
    else:
        try:
            text = content.decode('utf-8')
        except Exception:
            text = content.decode('latin-1')

    if not text.strip():
        raise HTTPException(400, "No text content found in file")

    doc_id = hashlib.md5(filename.encode()).hexdigest()[:12]
    (class_dir / f"{doc_id}.txt").write_text(text, encoding='utf-8')

    index = _get_index(user_id, class_id)
    index["documents"][doc_id] = {
        "title": doc_title,
        "filename": filename,
        "folder_id": folder_id,
    }
    _save_index(user_id, class_id, index)
    return {"doc_id": doc_id, "title": doc_title, "class_id": class_id, "folder_id": folder_id}


@router.delete("/documents/{doc_id}")
async def delete_document(doc_id: str, request: Request, class_id: str):
    user_id = _get_user_id(request)
    class_dir = _get_class_dir(user_id, class_id)
    doc_file = class_dir / f"{doc_id}.txt"
    if doc_file.exists():
        doc_file.unlink()
    index = _get_index(user_id, class_id)
    if doc_id in index["documents"]:
        del index["documents"][doc_id]
        _save_index(user_id, class_id, index)
    return {"deleted": True}


@router.patch("/documents/{doc_id}/move")
async def move_document(doc_id: str, request: Request):
    user_id = _get_user_id(request)
    body = await request.json()
    class_id = body.get("class_id")
    folder_id = body.get("folder_id")
    if not class_id:
        raise HTTPException(400, "class_id required")
    index = _get_index(user_id, class_id)
    if doc_id not in index["documents"]:
        raise HTTPException(404, "Document not found")
    index["documents"][doc_id]["folder_id"] = folder_id
    _save_index(user_id, class_id, index)
    return {"doc_id": doc_id, "folder_id": folder_id}


@router.get("/documents/{doc_id}/content")
async def get_document_content(doc_id: str, request: Request, class_id: str):
    user_id = _get_user_id(request)
    class_dir = _get_class_dir(user_id, class_id)
    doc_file = class_dir / f"{doc_id}.txt"
    if not doc_file.exists():
        raise HTTPException(404, "Document not found")
    content = doc_file.read_text(encoding="utf-8")
    index = _get_index(user_id, class_id)
    info = index["documents"].get(doc_id, {})
    return {"doc_id": doc_id, "title": info.get("title", doc_id), "content": content}


@router.get("/drive-files")
async def list_drive_files(request: Request):
    """List Google Docs and text files from Drive for import."""
    from app.services.drive import DriveService
    auth = request.headers.get("Authorization", "")
    token = auth[7:] if auth.startswith("Bearer ") else ""
    try:
        files = DriveService(token).list_files()
        return {"files": files}
    except Exception as e:
        raise HTTPException(400, f"Failed to list Drive files: {e}")


@router.post("/import-from-drive")
async def import_from_drive(
    request: Request,
    file_id: str = Form(...),
    class_id: str = Form("general"),
    folder_id: str = Form(None),
    title: str = Form(None),
):
    """Import a Google Doc or file from Drive into the file library."""
    from app.services.drive import DriveService
    import io
    auth = request.headers.get("Authorization", "")
    token = auth[7:] if auth.startswith("Bearer ") else ""
    user_id = _get_user_id(request)
    class_dir = _get_class_dir(user_id, class_id)
    service = DriveService(token)

    file_info = service.get_file(file_id)
    mime_type = file_info.get("mimeType", "")
    file_name = file_info.get("name", "untitled")
    doc_title = title or file_name

    if mime_type == "application/vnd.google-apps.document":
        text = service.export_as_text(file_id, mime_type)
    elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        content = service.download_file(file_id)
        from docx import Document as DocxDocument
        doc = DocxDocument(io.BytesIO(content))
        text = "\n".join(p.text for p in doc.paragraphs)
    elif mime_type == "application/pdf":
        content = service.download_file(file_id)
        from pypdf import PdfReader
        pdf = PdfReader(io.BytesIO(content))
        text = "\n".join(page.extract_text() or "" for page in pdf.pages)
    elif mime_type == "text/plain":
        text = service.download_file(file_id).decode("utf-8", errors="replace")
    else:
        raise HTTPException(400, f"Unsupported file type: {mime_type}")

    doc_id = hashlib.md5(f"{user_id}:{file_id}".encode()).hexdigest()[:16]
    (class_dir / f"{doc_id}.txt").write_text(text, encoding="utf-8")
    index = _get_index(user_id, class_id)
    index["documents"][doc_id] = {"title": doc_title, "filename": file_name, "source": "drive", "folder_id": folder_id}
    _save_index(user_id, class_id, index)
    return {"id": doc_id, "title": doc_title, "filename": file_name}


@router.patch("/documents/{doc_id}/rename")
async def rename_document(doc_id: str, request: Request):
    user_id = _get_user_id(request)
    body = await request.json()
    class_id = body.get("class_id")
    new_title = body.get("title", "").strip()
    if not class_id:
        raise HTTPException(400, "class_id required")
    if not new_title:
        raise HTTPException(400, "title required")
    index = _get_index(user_id, class_id)
    if doc_id not in index["documents"]:
        raise HTTPException(404, "Document not found")
    index["documents"][doc_id]["title"] = new_title
    _save_index(user_id, class_id, index)
    return {"doc_id": doc_id, "title": new_title}
