from fastapi import APIRouter, HTTPException, Request, UploadFile, File, Form
import json
import hashlib
from pathlib import Path
from app.services.token_store import get_user_id_from_token
from app.limiter import limiter

router = APIRouter()

# Simple file storage
DATA_DIR = Path(__file__).parent.parent.parent / "data" / "corpus"
DATA_DIR.mkdir(parents=True, exist_ok=True)


def get_user_id(request: Request) -> str:
    auth = request.headers.get("Authorization", "")
    if auth.startswith("Bearer "):
        token = auth[7:]
        return get_user_id_from_token(token)
    return "anonymous"


def get_class_dir(user_id: str, class_id: str) -> Path:
    class_dir = DATA_DIR / user_id / class_id
    class_dir.mkdir(parents=True, exist_ok=True)
    return class_dir


def get_index(user_id: str, class_id: str) -> dict:
    class_dir = get_class_dir(user_id, class_id)
    index_file = class_dir / "index.json"
    if index_file.exists():
        return json.loads(index_file.read_text())
    return {"documents": {}}


def save_index(user_id: str, class_id: str, index: dict):
    class_dir = get_class_dir(user_id, class_id)
    index_file = class_dir / "index.json"
    index_file.write_text(json.dumps(index, indent=2))


@router.get("/documents")
async def list_documents(request: Request, class_id: str):
    user_id = get_user_id(request)
    from app.services.linked_classes import LinkedClassService
    linked_service = LinkedClassService(user_id)
    linked_ids = linked_service.get_linked_class_ids(class_id)

    all_docs = []
    for cid in linked_ids:
        index = get_index(user_id, cid)
        for doc_id, info in index["documents"].items():
            doc = {"id": doc_id, **info, "source_class_id": cid}
            if cid != class_id:
                doc["shared"] = True
            all_docs.append(doc)
    return {"documents": all_docs, "class_id": class_id}


@router.post("/upload")
@limiter.limit("20/minute")
async def upload_document(
    request: Request,
    file: UploadFile = File(...),
    class_id: str = Form(...),
    title: str = Form(None)
):
    user_id = get_user_id(request)
    class_dir = get_class_dir(user_id, class_id)

    content = await file.read()
    filename = file.filename or "untitled"
    doc_title = title or filename

    # Handle different file types
    if filename.lower().endswith('.pdf'):
        try:
            from pypdf import PdfReader
            import io
            pdf = PdfReader(io.BytesIO(content))
            if pdf.is_encrypted:
                # Try decrypting with empty password (owner-locked but readable)
                result = pdf.decrypt("")
                if result == 0:
                    raise HTTPException(400, "This PDF is password-protected. Remove the password before uploading.")
            text = "\n".join(page.extract_text() or "" for page in pdf.pages)
        except HTTPException:
            raise
        except Exception as e:
            err = str(e)
            if "cryptography" in err.lower() or "AES" in err:
                raise HTTPException(400, "This PDF uses encryption that requires the 'cryptography' package. Please restart the backend and try again, or convert the PDF to an unencrypted version.")
            raise HTTPException(400, f"Failed to parse PDF: {e}")
    elif filename.lower().endswith('.docx'):
        try:
            from docx import Document
            import io
            doc = Document(io.BytesIO(content))
            text = "\n".join(para.text for para in doc.paragraphs)
        except Exception as e:
            raise HTTPException(400, f"Failed to parse Word document: {e}")
    else:
        try:
            text = content.decode('utf-8')
        except:
            text = content.decode('latin-1')

    if not text.strip():
        raise HTTPException(400, "No text content found in PDF - the PDF may be image-based")

    doc_id = hashlib.md5(filename.encode()).hexdigest()[:12]

    # Save document
    (class_dir / f"{doc_id}.txt").write_text(text, encoding='utf-8')

    # Update index
    index = get_index(user_id, class_id)
    index["documents"][doc_id] = {"title": doc_title, "filename": filename}
    save_index(user_id, class_id, index)

    return {"message": "Uploaded", "doc_id": doc_id, "title": doc_title, "class_id": class_id}


@router.post("/upload-text")
@limiter.limit("20/minute")
async def upload_text(
    request: Request,
    class_id: str = Form(...),
    title: str = Form(...),
    content: str = Form(...)
):
    user_id = get_user_id(request)
    class_dir = get_class_dir(user_id, class_id)

    if not content.strip():
        raise HTTPException(400, "Content cannot be empty")

    doc_id = hashlib.md5(title.encode()).hexdigest()[:12]

    # Save document
    (class_dir / f"{doc_id}.txt").write_text(content, encoding='utf-8')

    # Update index
    index = get_index(user_id, class_id)
    index["documents"][doc_id] = {"title": title, "type": "text"}
    save_index(user_id, class_id, index)

    return {"message": "Added", "doc_id": doc_id, "title": title, "class_id": class_id}


@router.delete("/documents/{doc_id}")
async def delete_document(doc_id: str, request: Request, class_id: str):
    user_id = get_user_id(request)
    class_dir = get_class_dir(user_id, class_id)

    doc_file = class_dir / f"{doc_id}.txt"
    if doc_file.exists():
        doc_file.unlink()

    index = get_index(user_id, class_id)
    if doc_id in index["documents"]:
        del index["documents"][doc_id]
        save_index(user_id, class_id, index)

    return {"deleted": True, "doc_id": doc_id, "class_id": class_id}


@router.get("/documents/{doc_id}/content")
async def get_document_content(doc_id: str, request: Request, class_id: str):
    user_id = get_user_id(request)
    class_dir = get_class_dir(user_id, class_id)
    index = get_index(user_id, class_id)

    if doc_id not in index["documents"]:
        raise HTTPException(404, "Document not found")

    doc_file = class_dir / f"{doc_id}.txt"
    if not doc_file.exists():
        raise HTTPException(404, "Document file not found")

    content = doc_file.read_text(encoding='utf-8')
    return {
        "doc_id": doc_id,
        "content": content,
        "metadata": index["documents"][doc_id]
    }


@router.get("/search")
async def search_corpus(request: Request, class_id: str, query: str, limit: int = 5):
    user_id = get_user_id(request)
    from app.services.linked_classes import LinkedClassService
    linked_service = LinkedClassService(user_id)
    linked_ids = linked_service.get_linked_class_ids(class_id)

    results = []
    query_lower = query.lower()

    for cid in linked_ids:
        cdir = get_class_dir(user_id, cid)
        idx_data = get_index(user_id, cid)
        for doc_id in idx_data["documents"]:
            doc_file = cdir / f"{doc_id}.txt"
            if doc_file.exists():
                text = doc_file.read_text(encoding='utf-8')
                if query_lower in text.lower():
                    pos = text.lower().find(query_lower)
                    snippet = text[max(0, pos-100):pos+300]
                    results.append({
                        "content": snippet,
                        "metadata": idx_data["documents"][doc_id],
                        "source_class_id": cid,
                    })

    return {"results": results[:limit], "class_id": class_id}


def get_token_from_request(request: Request) -> str:
    auth_header = request.headers.get("Authorization")
    if not auth_header or not auth_header.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Missing authorization")
    return auth_header.split(" ")[1]


@router.get("/drive-files")
async def list_drive_files(request: Request):
    """List Google Docs and text files from Drive that can be imported."""
    from app.services.drive import DriveService
    token = get_token_from_request(request)
    service = DriveService(token)
    try:
        files = service.list_docs_and_text_files()
        return {"files": files}
    except Exception as e:
        raise HTTPException(400, f"Failed to list Drive files: {e}")


@router.post("/import-from-drive")
async def import_from_drive(
    request: Request,
    file_id: str = Form(...),
    class_id: str = Form(...),
    title: str = Form(None)
):
    """Import a Google Doc or file from Drive into the corpus."""
    from app.services.drive import DriveService

    user_id = get_user_id(request)
    token = get_token_from_request(request)
    class_dir = get_class_dir(user_id, class_id)

    service = DriveService(token)

    try:
        # Get file info
        file_info = service.get_file(file_id)
        mime_type = file_info.get("mimeType", "")
        file_name = file_info.get("name", "untitled")
        doc_title = title or file_name

        # Extract text based on file type
        if mime_type == "application/vnd.google-apps.document":
            text = service.export_as_text(file_id, mime_type)
        elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            # Word doc - download and parse
            content = service.download_file(file_id)
            from docx import Document
            import io
            doc = Document(io.BytesIO(content))
            text = "\n".join(para.text for para in doc.paragraphs)
        elif mime_type == "application/pdf":
            content = service.download_file(file_id)
            from pypdf import PdfReader
            import io
            pdf = PdfReader(io.BytesIO(content))
            text = "\n".join(page.extract_text() or "" for page in pdf.pages)
        elif mime_type == "text/plain":
            content = service.download_file(file_id)
            text = content.decode('utf-8')
        else:
            raise HTTPException(400, f"Unsupported file type: {mime_type}")

        if not text.strip():
            raise HTTPException(400, "No text content found in file")

        doc_id = hashlib.md5(file_id.encode()).hexdigest()[:12]

        # Save document
        (class_dir / f"{doc_id}.txt").write_text(text, encoding='utf-8')

        # Update index
        index = get_index(user_id, class_id)
        index["documents"][doc_id] = {
            "title": doc_title,
            "source": "google_drive",
            "drive_file_id": file_id
        }
        save_index(user_id, class_id, index)

        return {
            "message": "Imported from Drive",
            "doc_id": doc_id,
            "title": doc_title,
            "class_id": class_id
        }

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(400, f"Failed to import: {e}")


@router.post("/import-from-url")
async def import_from_url(
    request: Request,
    url: str = Form(...),
    class_id: str = Form(...),
    title: str = Form(None)
):
    """Fetch a public URL (HTML page or PDF) and import its text into the corpus."""
    import httpx
    import re
    user_id = get_user_id(request)
    class_dir = get_class_dir(user_id, class_id)

    try:
        async with httpx.AsyncClient(follow_redirects=True, timeout=15) as client:
            resp = await client.get(url, headers={"User-Agent": "Mozilla/5.0"})
        resp.raise_for_status()
    except Exception as e:
        raise HTTPException(400, f"Could not fetch URL: {e}")

    content_type = resp.headers.get("content-type", "")

    if "application/pdf" in content_type:
        try:
            from pypdf import PdfReader
            import io
            pdf = PdfReader(io.BytesIO(resp.content))
            text = "\n".join(page.extract_text() or "" for page in pdf.pages)
        except Exception as e:
            raise HTTPException(400, f"Failed to parse PDF: {e}")
    elif "text/html" in content_type or "text/plain" in content_type:
        if "text/html" in content_type:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(resp.text, "html.parser")
            for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
                tag.decompose()
            text = soup.get_text(separator="\n")
            text = re.sub(r"\n{3,}", "\n\n", text).strip()
        else:
            text = resp.text
    else:
        raise HTTPException(400, f"Unsupported content type: {content_type}")

    if not text.strip():
        raise HTTPException(400, "No text content found at that URL")

    from urllib.parse import urlparse
    doc_title = title or urlparse(url).netloc + urlparse(url).path
    doc_id = hashlib.md5(url.encode()).hexdigest()[:12]

    (class_dir / f"{doc_id}.txt").write_text(text, encoding="utf-8")
    index = get_index(user_id, class_id)
    index["documents"][doc_id] = {"title": doc_title, "source": "url", "url": url}
    save_index(user_id, class_id, index)

    return {"message": "Imported from URL", "doc_id": doc_id, "title": doc_title, "class_id": class_id}
