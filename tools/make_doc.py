"""
make_doc.py — Convert a markdown file to a formatted Word (.docx) document.

Usage:
    python tools/make_doc.py legal/soulad-se-smernici-ai.md

Output: same path with .docx extension.

Supports:
  # H1  → Title style
  ## H2 → Heading 1
  ### H3 → Heading 2
  **bold**, plain text, bullet lists (- or •), tables (| col | col |), ✅ ⚠️ ❌
"""

import sys
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_run_with_markup(para, text):
    """Add a paragraph run, handling **bold** inline markup."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            para.add_run(part)


def style_heading(para, level):
    if level == 1:
        para.runs[0].font.size = Pt(22)
        para.runs[0].font.color.rgb = RGBColor(0x1a, 0x73, 0xe8)
        para.runs[0].bold = True
    elif level == 2:
        para.runs[0].font.size = Pt(14)
        para.runs[0].font.color.rgb = RGBColor(0x1a, 0x73, 0xe8)
        para.runs[0].bold = True
    elif level == 3:
        para.runs[0].font.size = Pt(12)
        para.runs[0].font.color.rgb = RGBColor(0x20, 0x20, 0x20)
        para.runs[0].bold = True


def convert(md_path: Path, docx_path: Path):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0

    while i < len(lines):
        line = lines[i]

        # Skip horizontal rules
        if re.match(r"^---+$", line.strip()):
            doc.add_paragraph()
            i += 1
            continue

        # H1
        if line.startswith("# ") and not line.startswith("## "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(line[2:])
            run.font.size = Pt(22)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1a, 0x73, 0xe8)
            i += 1
            continue

        # H2
        if line.startswith("## ") and not line.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(line[3:])
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1a, 0x73, 0xe8)
            # underline rule
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "1a73e8")
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # H3
        if line.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(line[4:])
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
            i += 1
            continue

        # Table
        if line.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].startswith("|"):
                row_line = lines[i]
                if re.match(r"^\|[-| :]+\|$", row_line.strip()):
                    i += 1
                    continue
                cells = [c.strip() for c in row_line.strip().strip("|").split("|")]
                rows.append(cells)
                i += 1
            if not rows:
                continue
            col_count = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=col_count)
            table.style = "Table Grid"
            for r_idx, row_data in enumerate(rows):
                row_obj = table.rows[r_idx]
                for c_idx, cell_text in enumerate(row_data):
                    if c_idx >= col_count:
                        break
                    cell = row_obj.cells[c_idx]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    add_run_with_markup(p, cell_text)
                    if r_idx == 0:
                        for run in p.runs:
                            run.bold = True
                        set_cell_bg(cell, "1a73e8")
                        for run in p.runs:
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            doc.add_paragraph()
            continue

        # Bullet list
        if re.match(r"^[-•*]\s", line):
            p = doc.add_paragraph(style="List Bullet")
            add_run_with_markup(p, line[2:])
            i += 1
            continue

        # Blockquote
        if line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            run = p.add_run(line[2:])
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # Bold metadata lines (e.g. **Prepared:** ...)
        if line.startswith("**") and ":**" in line:
            p = doc.add_paragraph()
            add_run_with_markup(p, line)
            i += 1
            continue

        # Empty line
        if line.strip() == "":
            doc.add_paragraph()
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        add_run_with_markup(p, line)
        i += 1

    doc.save(docx_path)
    print(f"Saved: {docx_path}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python tools/make_doc.py <input.md> [output.docx]")
        sys.exit(1)
    md = Path(sys.argv[1])
    docx = Path(sys.argv[2]) if len(sys.argv) > 2 else md.with_suffix(".docx")
    convert(md, docx)
